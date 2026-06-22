"""
document_parser.py
~~~~~~~~~~~~~~~~~~
Async-compatible document parser that supports PDF, DOCX, XLSX, TXT, and
common image formats.  Each parser returns a list of page dicts:

    [{"page_number": int, "text": str, "raw_images": list}, ...]

OCR is attempted on pages where text extraction yields too little content.
"""

import io
import logging
import os
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional heavy imports – fail gracefully so the service still starts even
# if a particular extra is not installed.
# ---------------------------------------------------------------------------
try:
    import pdfplumber  # type: ignore
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not installed – PDF parsing unavailable.")

try:
    import docx  # type: ignore
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed – DOCX parsing unavailable.")

try:
    import pandas as pd  # type: ignore
    _PANDAS_AVAILABLE = True
except ImportError:
    _PANDAS_AVAILABLE = False
    logger.warning("pandas not installed – XLSX parsing unavailable.")

try:
    from PIL import Image  # type: ignore
    _PIL_AVAILABLE = True
except ImportError:
    _PIL_AVAILABLE = False
    logger.warning("Pillow not installed – image parsing unavailable.")

try:
    import easyocr  # type: ignore
    _EASYOCR_AVAILABLE = True
except ImportError:
    _EASYOCR_AVAILABLE = False
    logger.warning("easyocr not installed – falling back to pytesseract for OCR.")

try:
    import pytesseract  # type: ignore
    _PYTESSERACT_AVAILABLE = True
except ImportError:
    _PYTESSERACT_AVAILABLE = False
    logger.warning("pytesseract not installed – OCR will be unavailable.")

# Minimum character count below which OCR is attempted for a PDF page
_PDF_OCR_THRESHOLD = 50

# Target word count per virtual page for DOCX / TXT chunking
_WORDS_PER_PAGE = 500


class DocumentParser:
    """Parse industrial documents into a uniform page-level representation."""

    # -----------------------------------------------------------------------
    # Public API
    # -----------------------------------------------------------------------

    def parse(self, file_path: str) -> List[Dict]:
        """Dispatch to the appropriate parser based on file extension.

        Args:
            file_path: Absolute or relative path to the document file.

        Returns:
            List of page dicts ``{"page_number": int, "text": str,
            "raw_images": list}``.  Returns an empty list if the file is
            unparseable or an unrecognised type.
        """
        path = Path(file_path)
        ext = path.suffix.lower().lstrip(".")

        try:
            if ext == "pdf":
                return self._parse_pdf(file_path)
            elif ext in {"docx", "doc"}:
                return self._parse_docx(file_path)
            elif ext in {"xlsx", "xls", "csv"}:
                return self._parse_xlsx(file_path)
            elif ext == "txt":
                return self._parse_txt(file_path)
            elif ext in {"png", "jpg", "jpeg", "bmp", "tiff", "tif", "webp"}:
                return self._parse_image(file_path)
            else:
                logger.warning("Unsupported file extension: .%s – skipping %s", ext, file_path)
                return []
        except Exception as exc:  # noqa: BLE001
            logger.error("Failed to parse %s: %s", file_path, exc, exc_info=True)
            return []

    # -----------------------------------------------------------------------
    # Format-specific parsers
    # -----------------------------------------------------------------------

    def _parse_pdf(self, file_path: str) -> List[Dict]:
        """Extract text from each PDF page using pdfplumber.

        If a page yields fewer than :data:`_PDF_OCR_THRESHOLD` characters,
        the page is rendered to an image and OCR is attempted.

        Args:
            file_path: Path to the PDF file.

        Returns:
            List of page dicts.
        """
        if not _PDFPLUMBER_AVAILABLE:
            logger.error("pdfplumber is required to parse PDF files.")
            return []

        pages: List[Dict] = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text: str = page.extract_text() or ""

                    if len(text.strip()) < _PDF_OCR_THRESHOLD:
                        logger.debug(
                            "Page %d has sparse text (%d chars) – attempting OCR.",
                            page_num,
                            len(text.strip()),
                        )
                        try:
                            # Render page to an image for OCR
                            page_image = page.to_image(resolution=200)
                            img_bytes_io = io.BytesIO()
                            page_image.original.save(img_bytes_io, format="PNG")
                            ocr_text = self._ocr_text(img_bytes_io.getvalue())
                            if ocr_text:
                                text = ocr_text
                        except Exception as ocr_exc:  # noqa: BLE001
                            logger.warning(
                                "OCR failed for page %d of %s: %s",
                                page_num,
                                file_path,
                                ocr_exc,
                            )

                    pages.append(
                        {
                            "page_number": page_num,
                            "text": text.strip(),
                            "raw_images": [],
                        }
                    )
        except Exception as exc:  # noqa: BLE001
            logger.error("pdfplumber failed on %s: %s", file_path, exc, exc_info=True)

        return pages

    def _parse_docx(self, file_path: str) -> List[Dict]:
        """Extract text from a DOCX file and group it into virtual pages.

        Paragraphs and table cells are concatenated and then split into
        virtual pages of approximately :data:`_WORDS_PER_PAGE` words each.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            List of virtual-page dicts.
        """
        if not _DOCX_AVAILABLE:
            logger.error("python-docx is required to parse DOCX files.")
            return []

        all_text_parts: List[str] = []

        try:
            document = docx.Document(file_path)

            # Extract paragraphs
            for para in document.paragraphs:
                stripped = para.text.strip()
                if stripped:
                    all_text_parts.append(stripped)

            # Extract tables
            for table in document.tables:
                for row in table.rows:
                    row_parts: List[str] = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_parts.append(cell_text)
                    if row_parts:
                        all_text_parts.append(" | ".join(row_parts))

        except Exception as exc:  # noqa: BLE001
            logger.error("python-docx failed on %s: %s", file_path, exc, exc_info=True)
            return []

        return self._split_into_virtual_pages(" ".join(all_text_parts))

    def _parse_xlsx(self, file_path: str) -> List[Dict]:
        """Convert each sheet of an Excel / CSV file into a text representation.

        Args:
            file_path: Path to the XLSX / XLS / CSV file.

        Returns:
            List of page dicts, one per sheet.
        """
        if not _PANDAS_AVAILABLE:
            logger.error("pandas is required to parse XLSX / CSV files.")
            return []

        pages: List[Dict] = []
        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".csv":
                # Treat a CSV as a single-sheet workbook
                df = pd.read_csv(file_path, dtype=str).fillna("")
                sheet_text = f"Sheet: {Path(file_path).stem}\n\n{df.to_string(index=False)}"
                pages.append({"page_number": 1, "text": sheet_text, "raw_images": []})
            else:
                xl = pd.ExcelFile(file_path)
                for page_num, sheet_name in enumerate(xl.sheet_names, start=1):
                    df = xl.parse(sheet_name, dtype=str).fillna("")
                    sheet_text = f"Sheet: {sheet_name}\n\n{df.to_string(index=False)}"
                    pages.append(
                        {
                            "page_number": page_num,
                            "text": sheet_text,
                            "raw_images": [],
                        }
                    )
        except Exception as exc:  # noqa: BLE001
            logger.error("pandas failed on %s: %s", file_path, exc, exc_info=True)

        return pages

    def _parse_txt(self, file_path: str) -> List[Dict]:
        """Read a plain-text file and split it into virtual pages.

        Args:
            file_path: Path to the TXT file.

        Returns:
            List of virtual-page dicts.
        """
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
                raw_text = fh.read()
        except Exception as exc:  # noqa: BLE001
            logger.error("Failed to read text file %s: %s", file_path, exc, exc_info=True)
            return []

        return self._split_into_virtual_pages(raw_text)

    def _parse_image(self, file_path: str) -> List[Dict]:
        """Run OCR on an image file to extract text.

        Args:
            file_path: Path to the image file.

        Returns:
            Single-element list with the OCR result, or empty list on failure.
        """
        try:
            with open(file_path, "rb") as fh:
                image_bytes = fh.read()
            text = self._ocr_text(image_bytes)
            return [{"page_number": 1, "text": text.strip(), "raw_images": []}]
        except Exception as exc:  # noqa: BLE001
            logger.error("Image parsing failed for %s: %s", file_path, exc, exc_info=True)
            return []

    # -----------------------------------------------------------------------
    # OCR helper
    # -----------------------------------------------------------------------

    def _ocr_text(self, image_bytes: bytes) -> str:
        """Extract text from raw image bytes using OCR.

        EasyOCR is tried first; if unavailable or it raises an exception,
        pytesseract is used as a fallback.

        Args:
            image_bytes: Raw bytes of the image (any format supported by Pillow).

        Returns:
            Extracted text string (may be empty if OCR finds nothing).
        """
        # ---- EasyOCR -------------------------------------------------------
        if _EASYOCR_AVAILABLE:
            try:
                reader = easyocr.Reader(["en"], gpu=False, verbose=False)
                results = reader.readtext(image_bytes, detail=0, paragraph=True)
                return " ".join(results)
            except Exception as exc:  # noqa: BLE001
                logger.warning("EasyOCR failed – falling back to pytesseract: %s", exc)

        # ---- pytesseract fallback ------------------------------------------
        if _PYTESSERACT_AVAILABLE and _PIL_AVAILABLE:
            try:
                image = Image.open(io.BytesIO(image_bytes))
                return pytesseract.image_to_string(image)
            except Exception as exc:  # noqa: BLE001
                logger.error("pytesseract OCR also failed: %s", exc, exc_info=True)
                return ""

        logger.error(
            "Neither EasyOCR nor pytesseract is available – cannot perform OCR."
        )
        return ""

    # -----------------------------------------------------------------------
    # Internal helpers
    # -----------------------------------------------------------------------

    def _split_into_virtual_pages(self, text: str) -> List[Dict]:
        """Split a long text string into virtual pages of ~:data:`_WORDS_PER_PAGE` words.

        Args:
            text: The full document text as a single string.

        Returns:
            List of virtual-page dicts with incrementing page numbers.
        """
        words = text.split()
        if not words:
            return []

        pages: List[Dict] = []
        page_num = 1

        for start in range(0, len(words), _WORDS_PER_PAGE):
            chunk_words = words[start : start + _WORDS_PER_PAGE]
            pages.append(
                {
                    "page_number": page_num,
                    "text": " ".join(chunk_words),
                    "raw_images": [],
                }
            )
            page_num += 1

        return pages
